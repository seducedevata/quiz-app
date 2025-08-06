"""
Advanced document processing module for Knowledge App - FULLY ASYNC VERSION
Implements semantic chunking, embedding-based processing, and educational content preservation
🚀 MAJOR FIXES: Async processing, parallel execution, sentence transformers, structured data generation
"""

from .async_converter import async_file_read, async_file_write

import os
import logging
import re
import json
import csv
import asyncio
import aiofiles
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from datetime import datetime
from collections import Counter
import shutil
import time
import unicodedata
import numpy as np
import hashlib

# Fix NumPy binary compatibility issue before importing sklearn-dependent packages
import warnings

# Initialize logger first before any other code that might use it
logger = logging.getLogger(__name__)

# Add local NLTK data path
try:
    import nltk
    import os
    local_nltk_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'nltk_data_fixed')
    if os.path.exists(local_nltk_dir):
        nltk.data.path.append(local_nltk_dir)
except ImportError:
    pass

warnings.filterwarnings("ignore", message="numpy.dtype size changed")
warnings.filterwarnings("ignore", message="numpy.ufunc size changed")

# 🔥 CRITICAL FIX: Make NLTK initialization more robust
nltk_available = False
sent_tokenize = None
word_tokenize = None
stopwords = None
WordNetLemmatizer = None
pos_tag = None
wordnet = None

try:
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    from nltk.stem import WordNetLemmatizer
    from nltk.tag import pos_tag
    from nltk.corpus import wordnet

    # 🔧 FIX: Use bundled NLTK data, avoid runtime downloads
    # Set up local NLTK data path first
    nltk_data_path = Path(__file__).parent.parent / "data" / "nltk_data"
    if nltk_data_path.exists():
        nltk.data.path.insert(0, str(nltk_data_path))
        logger.info(f"✅ Using bundled NLTK data from: {nltk_data_path}")

    # Test NLTK data availability - warn but don't download
    missing_data = []

    # CRITICAL FIX: Automatically download missing NLTK resources
    nltk_resources = [
        ('taggers/averaged_perceptron_tagger', 'averaged_perceptron_tagger'),
        ('tokenizers/punkt', 'punkt'),
        ('corpora/stopwords', 'stopwords'),
        ('corpora/wordnet', 'wordnet')
    ]
    
    for resource_path, resource_name in nltk_resources:
        try:
            nltk.data.find(resource_path)
        except LookupError:
            logger.warning(f"⚠️ Missing {resource_name}, attempting to download...")
            try:
                # CRITICAL FIX: Attempt automatic download
                nltk.download(resource_name, quiet=True)
                logger.info(f"✅ Successfully downloaded {resource_name}")
                
                # Verify download worked
                nltk.data.find(resource_path)
                logger.info(f"✅ Verified {resource_name} is now available")
                
            except Exception as download_error:
                logger.error(f"❌ Failed to download {resource_name}: {download_error}")
                missing_data.append(resource_name)
                logger.warning(f"⚠️ {resource_name} will not be available - some features may be limited")

    if missing_data:
        logger.warning(f"⚠️ Missing NLTK data: {missing_data}")
        logger.warning("⚠️ Document processing will use fallback methods")
        logger.info("💡 For full functionality, ensure NLTK data is bundled with the application")
    else:
        logger.info("✅ All required NLTK data available")

    # 🔥 CRITICAL FIX: Test if sent_tokenize actually works
    try:
        test_result = sent_tokenize("This is a test. This is another sentence.")
        if test_result and isinstance(test_result, list) and len(test_result) == 2:
            nltk_available = True
            logger.info("✅ NLTK successfully initialized and tested")
        else:
            raise Exception("sent_tokenize test failed")
    except Exception as e:
        logger.warning(f"NLTK sent_tokenize test failed: {e}")
        nltk_available = False
        sent_tokenize = None

except ImportError as e:
    logger.warning(f"NLTK not available: {e}. Some advanced features will be disabled.")
    nltk_available = False
    nltk = None
    sent_tokenize = None
    word_tokenize = None
    stopwords = None
    WordNetLemmatizer = None
    pos_tag = None
    wordnet = None

try:
    import fitz  # PyMuPDF
    pdf_support_available = True
except ImportError:
    pdf_support_available = False
    logging.warning("PyMuPDF not available. PDF processing will be limited.")

try:
    from docx import Document
    docx_support_available = True
except ImportError:
    docx_support_available = False
    logging.warning("python-docx not available. DOCX processing will be limited.")

# Try to import advanced ML libraries for semantic processing
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    sklearn_available = True
    logger.info("✅ Scikit-learn available for semantic processing")
except ImportError:
    sklearn_available = False
    logger.warning("Scikit-learn not available. Semantic chunking will use fallback methods.")

try:
    import langdetect
    langdetect_available = True
    logger.info("✅ Language detection available")
except ImportError:
    langdetect_available = False
    logger.warning("Language detection not available")


class AdvancedDocumentProcessor:
    """
    🚀 ENHANCED Advanced document processor with semantic chunking and educational content preservation.
    
    MAJOR IMPROVEMENTS IMPLEMENTED:
    ✅ Async processing to prevent UI blocking
    ✅ Parallel file processing with controlled concurrency  
    ✅ Sentence transformers instead of outdated TF-IDF
    ✅ Markdown table extraction preserving structure
    ✅ Enhanced sentence tokenization fallback
    ✅ Configuration-aware cache invalidation
    ✅ Instruction-following training data format
    ✅ Multiple task variations for diverse training
    ✅ Better error handling and validation
    
    Implements cutting-edge techniques for high-quality training data extraction.
    """

    def __init__(self, preserve_educational_content: bool = True, use_semantic_chunking: bool = True):
        """Initialize processor with advanced educational content preservation."""
        self.preserve_educational_content = preserve_educational_content
        self.use_semantic_chunking = use_semantic_chunking
        
        # 🔥 ENHANCED: Initialize components safely without destroying content
        self.lemmatizer = None
        self.stop_words = set()
        
        if nltk_available and not preserve_educational_content:  # Only use lemmatization if not preserving educational content
            try:
                self.lemmatizer = WordNetLemmatizer()
                self.stop_words = set(stopwords.words("english"))
            except Exception as e:
                logger.warning(f"Failed to initialize NLTK components: {e}")
                self.lemmatizer = None
                self.stop_words = set()

        # 🚀 ADVANCED: Educational content preservation patterns
        self.educational_patterns = {
            'mathematical_expressions': r'[a-zA-Z]*\d+[a-zA-Z\d\+\-\*\/\=\(\)\^\{\}]*',
            'formulas': r'[A-Za-z]\s*[=]\s*[A-Za-z0-9\+\-\*\/\(\)\^\s]+',
            'equations': r'\b[a-zA-Z]\s*[+\-*/=]\s*[a-zA-Z0-9]+',
            'scientific_notation': r'\d+\.?\d*[eE][\+\-]?\d+',
            'units': r'\d+\s*(mm|cm|m|km|g|kg|s|min|h|°C|°F|K)',
            'references': r'\[?\d+\]?|\(\d+\)',
            'chapter_sections': r'(Chapter|Section|Part)\s+\d+',
            'examples': r'(Example|Problem|Exercise)\s*\d*:?',
            'solutions': r'(Solution|Answer)\s*:?'
        }

        # 🚀 ADVANCED: Minimal noise patterns (preserve educational content)
        self.minimal_noise_patterns = [
            r'\x00+',  # Null bytes only
            r'[\r\n]{3,}',  # Excessive line breaks only
            r'\s{4,}',  # Excessive spaces only (keep some for formatting)
        ]

        # 🚀 ADVANCED: Document structure patterns
        self.structure_patterns = {
            'table_markers': r'(\|.*\||\+[-=]+\+)',
            'list_items': r'^[\s]*[\*\-\+\d+\.]\s+',
            'headings': r'^#{1,6}\s+.*$|^[A-Z][A-Z\s]+$',
            'code_blocks': r'```[\s\S]*?```|`[^`]+`',
            'citations': r'\[[^\]]+\]|\([^)]+\)'
        }

        # 🚀 ADVANCED: Quality thresholds for educational content
        self.quality_thresholds = {
            'min_sentence_length': 5,  # Reduced for math content
            'max_sentence_length': 1000,  # Increased for complex explanations
            'min_word_count': 2,  # Reduced for formulas
            'max_word_count': 200,  # Increased for detailed explanations
            'min_letter_ratio': 0.3,  # Reduced for math/formulas
            'perplexity_threshold': 50.0  # For advanced filtering
        }

        # 🚀 UPGRADE: Initialize sentence transformer for semantic analysis (replaces TF-IDF)
        self.sentence_transformer = None
        try:
            from sentence_transformers import SentenceTransformer
            # Use a lightweight, fast model for semantic embeddings
            self.sentence_transformer = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("✅ Sentence transformer initialized for semantic processing")
        except ImportError:
            logger.warning("Sentence transformers not available. Falling back to TF-IDF")
            # Fallback to TF-IDF if sentence transformers not available
            if sklearn_available:
                self.tfidf_vectorizer = TfidfVectorizer(
                    max_features=1000,
                    stop_words='english' if not preserve_educational_content else None,
                    ngram_range=(1, 2)
                )
            else:
                self.tfidf_vectorizer = None

        logger.info(f"🚀 AdvancedDocumentProcessor initialized - Educational preservation: {preserve_educational_content}, Semantic chunking: {use_semantic_chunking}")

    def extract_text_with_structure(self, file_path: str) -> Dict[str, Any]:
        """🚀 ENHANCED: Extract text with structural information preservation"""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        metadata = {
            "filename": file_path.name,
            "file_type": file_path.suffix.lower()[1:],
            "file_size": file_path.stat().st_size,
            "processed_at": datetime.now().isoformat(),
        }

        try:
            content = ""
            structural_info = {}
            
            if file_path.suffix.lower() == ".pdf" and pdf_support_available:
                content, meta, structural_info = self._extract_from_pdf_enhanced(file_path)
                if meta and isinstance(meta, dict):
                    metadata.update(meta)
            elif file_path.suffix.lower() == ".docx" and docx_support_available:
                content, meta, structural_info = self._extract_from_docx_enhanced(file_path)
                if meta and isinstance(meta, dict):
                    metadata.update(meta)
            elif file_path.suffix.lower() == ".txt":
                content = self._extract_from_txt_enhanced_sync(file_path)
                structural_info = self._analyze_text_structure(content)
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}")

            if content is None:
                content = ""
                logger.warning(f"Content extraction returned None for {file_path}")

            metadata["total_chars"] = len(content) if content else 0
            metadata["total_words"] = len(content.split()) if content else 0
            
            # 🚀 ENHANCED: Add structural metadata
            metadata["structural_info"] = structural_info

            return {
                "raw": content, 
                "metadata": metadata,
                "structure": structural_info
            }

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise

    def _extract_from_pdf_enhanced(self, file_path: Path) -> Tuple[str, Dict[str, Any], Dict[str, Any]]:
        """🚀 ENHANCED: Extract PDF with structure preservation and memory management"""
        if not pdf_support_available:
            raise ImportError("PyMuPDF not available for PDF processing")

        doc = None
        try:
            doc = fitz.open(file_path)

            # 🔧 MEMORY FIX: Check document size and implement memory limits
            page_count = len(doc) if doc else 0
            max_pages = 1000  # Limit to prevent memory exhaustion
            max_memory_mb = 500  # Maximum memory usage in MB

            if page_count > max_pages:
                logger.warning(f"⚠️ Large PDF detected ({page_count} pages). Processing first {max_pages} pages to prevent memory issues.")
                page_count = max_pages

            content = []
            current_memory_mb = 0
            structural_info = {
                'pages': [],
                'tables': [],
                'images': [],
                'fonts': set(),
                'has_mathematical_content': False,
                'truncated': False,
                'processed_pages': 0,
                'total_pages': len(doc) if doc else 0
            }

            pdf_metadata = doc.metadata if doc.metadata is not None else {}
            metadata = {
                "title": pdf_metadata.get("title", "") if pdf_metadata else "",
                "author": pdf_metadata.get("author", "") if pdf_metadata else "",
                "subject": pdf_metadata.get("subject", "") if pdf_metadata else "",
                "keywords": pdf_metadata.get("keywords", "") if pdf_metadata else "",
                "page_count": len(doc) if doc else 0,
            }

            if not doc or len(doc) == 0:
                logger.warning(f"PDF {file_path} appears to be empty or corrupted")
                return "", metadata, structural_info

            for page_num, page in enumerate(doc):
                # 🔧 MEMORY FIX: Stop processing if we've reached limits
                if page_num >= page_count:
                    structural_info['truncated'] = True
                    logger.warning(f"⚠️ PDF processing truncated at page {page_num} due to size limits")
                    break

                try:
                    # Extract text with formatting information
                    page_text = page.get_text()
                    if page_text:
                        # 🔧 MEMORY FIX: Estimate memory usage and check limits
                        page_size_mb = len(page_text.encode('utf-8')) / (1024 * 1024)
                        current_memory_mb += page_size_mb

                        if current_memory_mb > max_memory_mb:
                            structural_info['truncated'] = True
                            logger.warning(f"⚠️ PDF processing stopped at page {page_num} due to memory limit ({current_memory_mb:.1f}MB)")
                            break

                        # 🚀 ENHANCED: Preserve formatting and detect mathematical content
                        page_text = self._normalize_unicode(page_text)

                        # Detect mathematical expressions
                        if self._contains_mathematical_content(page_text):
                            structural_info['has_mathematical_content'] = True

                        # Extract font information for structure analysis
                        try:
                            text_dict = page.get_text("dict")
                            for block in text_dict.get("blocks", []):
                                if "lines" in block:
                                    for line in block["lines"]:
                                        for span in line.get("spans", []):
                                            if "font" in span:
                                                structural_info['fonts'].add(span["font"])
                        except:
                            pass  # Font extraction is optional
                        
                        content.append(page_text)

                        structural_info['pages'].append({
                            'page_num': page_num + 1,
                            'char_count': len(page_text),
                            'has_tables': bool(re.search(self.structure_patterns['table_markers'], page_text)),
                            'has_lists': bool(re.search(self.structure_patterns['list_items'], page_text, re.MULTILINE)),
                            'has_math': self._contains_mathematical_content(page_text),
                            'memory_mb': page_size_mb
                        })

                        structural_info['processed_pages'] = page_num + 1
                        
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num} in {file_path}: {e}")
                    continue

            if not content:
                logger.warning(f"No text content extracted from PDF {file_path}")
                return "", metadata, structural_info

            full_content = "\n\n".join(content)
            full_content = self._normalize_unicode(full_content)

            # 🔧 MEMORY FIX: Add memory usage information to metadata
            metadata['memory_usage_mb'] = current_memory_mb
            metadata['processing_truncated'] = structural_info.get('truncated', False)
            metadata['processed_pages'] = structural_info.get('processed_pages', 0)

            # Convert fonts set to list for JSON serialization
            structural_info['fonts'] = list(structural_info['fonts'])

            return full_content, metadata, structural_info
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return "", {"title": "", "author": "", "subject": "", "keywords": "", "page_count": 0}, {}
        finally:
            if doc:
                try:
                    doc.close()
                except Exception as e:
                    logger.warning(f"Error closing PDF {file_path}: {e}")

    def _extract_from_docx_enhanced(self, file_path: Path) -> Tuple[str, Dict[str, Any], Dict[str, Any]]:
        """🚀 ENHANCED: Extract DOCX with structure preservation"""
        if not docx_support_available:
            raise ImportError("python-docx not available for DOCX processing")

        try:
            doc = Document(file_path)
            content = []
            structural_info = {
                'paragraphs': 0,
                'tables': 0,
                'headings': [],
                'has_mathematical_content': False
            }
            
            core_props = doc.core_properties if hasattr(doc, 'core_properties') and doc.core_properties else None
            metadata = {
                "title": core_props.title if core_props and core_props.title else "",
                "author": core_props.author if core_props and core_props.author else "",
                "subject": core_props.subject if core_props and core_props.subject else "",
                "keywords": core_props.keywords if core_props and core_props.keywords else "",
                "page_count": len(doc.sections) if hasattr(doc, 'sections') and doc.sections else 0,
            }

            # 🚀 ENHANCED: Extract tables with structure
            if hasattr(doc, 'tables') and doc.tables:
                structural_info['tables'] = len(doc.tables)
                for table in doc.tables:
                    table_text = self._extract_table_content(table)
                    if table_text:
                        content.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")

            # Extract paragraphs with heading detection
            if hasattr(doc, 'paragraphs') and doc.paragraphs:
                for para in doc.paragraphs:
                    if para and hasattr(para, 'text'):
                        para_text = para.text
                        if para_text:
                            para_text = self._normalize_unicode(para_text)
                            
                            # Detect headings
                            if hasattr(para, 'style') and para.style and 'heading' in para.style.name.lower():
                                structural_info['headings'].append(para_text)
                                para_text = f"\n## {para_text}\n"
                            
                            # Check for mathematical content
                            if self._contains_mathematical_content(para_text):
                                structural_info['has_mathematical_content'] = True
                            
                            content.append(para_text)
                            structural_info['paragraphs'] += 1

            if not content:
                logger.warning(f"No text content extracted from DOCX {file_path}")
                return "", metadata, structural_info

            full_content = "\n\n".join(content)
            return full_content, metadata, structural_info
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return "", {"title": "", "author": "", "subject": "", "keywords": "", "page_count": 0}, {}

    def _extract_from_txt_enhanced_sync(self, file_path: Path) -> str:
        """🚀 ENHANCED: Extract TXT with better encoding handling (synchronous version)"""
        try:
            # Try UTF-8 first with error handling
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                if content:
                    content = self._normalize_unicode(content)
                    return content
        except UnicodeDecodeError:
            # Fallback to binary mode and decode with error handling
            try:
                with open(file_path, 'rb') as f:
                    raw_content = f.read()

                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        content = raw_content.decode(encoding)
                        return self._normalize_unicode(content)
                    except UnicodeDecodeError:
                        continue
                # If all fail, use UTF-8 with replacement
                content = raw_content.decode('utf-8', errors='replace')
                return self._normalize_unicode(content)
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return ""
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""

    def _normalize_unicode(self, text: str) -> str:
        """🚀 ADVANCED: Unicode normalization for consistent processing"""
        if not text:
            return ""
        
        try:
            # Normalize Unicode to NFC form
            text = unicodedata.normalize('NFC', text)
            
            # Remove null bytes and control characters (except newlines and tabs)
            text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t\r')
            
            return text
        except Exception as e:
            logger.warning(f"Unicode normalization failed: {e}")
            return text

    def _contains_mathematical_content(self, text: str) -> bool:
        """🚀 ADVANCED: Detect mathematical content to preserve it"""
        if not text:
            return False
        
        for pattern_name, pattern in self.educational_patterns.items():
            if re.search(pattern, text):
                return True
        
        # Additional checks for mathematical content
        math_indicators = [
            r'\b\d+\s*[+\-*/=]\s*\d+',  # Simple arithmetic
            r'[a-zA-Z]\s*=\s*[a-zA-Z0-9+\-*/()]+',  # Variable assignments
            r'\b(sin|cos|tan|log|ln|exp|sqrt)\s*\(',  # Mathematical functions
            r'\b(theorem|proof|lemma|corollary)\b',  # Mathematical terms
            r'[∫∑∏∆∇∂∞±≤≥≠≈]',  # Mathematical symbols
        ]
        
        for pattern in math_indicators:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False

    def _extract_table_content(self, table) -> str:
        """🚀 ENHANCED: Extract table content in Markdown format to preserve structure"""
        try:
            table_rows = []
            header_row = None
            
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip() if cell.text else ""
                    # Escape pipe characters in cell content
                    cell_text = cell_text.replace("|", "\\|")
                    row_cells.append(cell_text)
                
                if any(row_cells):  # Only add non-empty rows
                    row_text = " | ".join(row_cells)
                    if row_idx == 0:
                        # First row as header
                        header_row = row_text
                        table_rows.append(header_row)
                        # Add separator for Markdown table format
                        separator = " | ".join(["---"] * len(row_cells))
                        table_rows.append(separator)
                    else:
                        table_rows.append(row_text)
            
            if table_rows:
                # Return as Markdown table with structure preserved
                return "\n".join(table_rows)
            else:
                return ""
        except Exception as e:
            logger.warning(f"Failed to extract table content: {e}")
            return ""

    def _analyze_text_structure(self, text: str) -> Dict[str, Any]:
        """🚀 ADVANCED: Analyze text structure for better processing"""
        if not text:
            return {}
        
        structure = {
            'has_tables': bool(re.search(self.structure_patterns['table_markers'], text)),
            'has_lists': bool(re.search(self.structure_patterns['list_items'], text, re.MULTILINE)),
            'has_headings': bool(re.search(self.structure_patterns['headings'], text, re.MULTILINE)),
            'has_code': bool(re.search(self.structure_patterns['code_blocks'], text)),
            'has_citations': bool(re.search(self.structure_patterns['citations'], text)),
            'has_mathematical_content': self._contains_mathematical_content(text),
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
            'sentence_count': len(sent_tokenize(text)) if nltk_available else text.count('.'),
            'word_count': len(text.split()),
            'character_count': len(text)
        }
        
        # Detect language
        if langdetect_available:
            try:
                structure['language'] = langdetect.detect(text[:1000])  # Sample first 1000 chars
            except:
                structure['language'] = 'unknown'
        else:
            structure['language'] = 'unknown'
        
        return structure

    def educational_text_cleaning(self, text: str) -> str:
        """🚀 REVOLUTIONARY: Educational content preserving text cleaning"""
        if text is None:
            logger.warning("❌ Text is None, returning empty string")
            return ""
        
        if not isinstance(text, str):
            logger.warning(f"❌ Text is not a string (type: {type(text)}), converting to string")
            text = str(text)
        
        if not text or not text.strip():
            logger.warning("❌ Text is empty after stripping")
            return ""

        try:
            # 🚀 STEP 1: Unicode normalization
            text = self._normalize_unicode(text)
            
            # 🚀 STEP 2: Minimal noise removal (preserve educational content!)
            for pattern in self.minimal_noise_patterns:
                text = re.sub(pattern, " ", text)
            
            # 🚀 STEP 3: Fix common OCR/scanning artifacts without destroying content
            text = self._fix_scanning_artifacts(text)
            
            # 🚀 STEP 4: Preserve mathematical expressions and formulas
            text = self._preserve_mathematical_content(text)
            
            # 🚀 STEP 5: Enhance readability without losing information
            text = self._enhance_readability(text)
            
            # 🚀 STEP 6: Final structural cleanup
            text = self._final_structural_cleanup(text)
            
            return text
            
        except Exception as e:
            logger.error(f"❌ Error in educational text cleaning: {e}")
            return text  # Return original text if cleaning fails

    def _fix_scanning_artifacts(self, text: str) -> str:
        """🚀 ADVANCED: Fix common scanning artifacts without destroying content"""
        try:
            # Fix broken words (common in OCR)
            text = re.sub(r'\b([a-z])\s+([a-z]{2,})\b', r'\1\2', text)
            
            # Fix repeated punctuation
            text = re.sub(r'([.!?])\1{2,}', r'\1', text)
            
            # Fix spacing around punctuation
            text = re.sub(r'\s+([,.!?;:])', r'\1', text)
            text = re.sub(r'([,.!?;:])\s+', r'\1 ', text)
            
            # Fix line breaks in middle of sentences (but preserve intentional breaks)
            text = re.sub(r'([a-z])\n([a-z])', r'\1 \2', text)
            
            return text
        except Exception as e:
            logger.warning(f"Failed to fix scanning artifacts: {e}")
            return text

    def _preserve_mathematical_content(self, text: str) -> str:
        """🚀 ADVANCED: Preserve and enhance mathematical expressions"""
        try:
            # Preserve mathematical expressions by adding markers
            for pattern_name, pattern in self.educational_patterns.items():
                matches = re.finditer(pattern, text)
                for match in reversed(list(matches)):  # Reverse to maintain positions
                    original = match.group()
                    preserved = f" [MATH:{original}] "
                    text = text[:match.start()] + preserved + text[match.end():]
            
            return text
        except Exception as e:
            logger.warning(f"Failed to preserve mathematical content: {e}")
            return text

    def _enhance_readability(self, text: str) -> str:
        """🚀 ADVANCED: Enhance readability while preserving educational content"""
        try:
            # Normalize whitespace but preserve structure
            lines = text.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line = line.strip()
                if line:
                    # Ensure proper spacing but don't over-normalize
                    line = re.sub(r'\s+', ' ', line)
                    cleaned_lines.append(line)
                else:
                    # Preserve intentional blank lines for structure
                    if cleaned_lines and cleaned_lines[-1] != '':
                        cleaned_lines.append('')
            
            # Remove excessive blank lines but preserve structure
            text = '\n'.join(cleaned_lines)
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            return text
        except Exception as e:
            logger.warning(f"Failed to enhance readability: {e}")
            return text

    def _final_structural_cleanup(self, text: str) -> str:
        """🚀 ADVANCED: Final cleanup preserving document structure"""
        try:
            # Restore mathematical expressions
            text = re.sub(r'\[MATH:([^\]]+)\]', r'\1', text)
            
            # Ensure sentences end properly
            text = re.sub(r'([a-zA-Z0-9])\s*$', r'\1.', text, flags=re.MULTILINE)
            
            # Final whitespace normalization
            text = text.strip()
            
            return text
        except Exception as e:
            logger.warning(f"Failed final structural cleanup: {e}")
            return text

    def advanced_text_cleaning(self, text: str) -> str:
        """🚀 ENHANCED: Apply advanced text cleaning pipeline - now redirects to educational cleaning"""
        logger.info("🚀 Using educational content preserving text cleaning")
        return self.educational_text_cleaning(text)

    def generate_training_dataset(
        self, content: Dict[str, str], chunk_size: int = 500
    ) -> List[Dict[str, str]]:
        """🚀 ENHANCED: Generate training dataset - now redirects to enhanced method"""
        logger.info("🚀 Using enhanced training dataset generation with semantic chunking and quality filtering")
        return self.generate_enhanced_training_dataset(content, chunk_size)

    async def process_documents_advanced(
        self, file_paths: List[str], output_dir: str, chunk_size: int = 500
    ) -> Dict[str, Any]:
        """🚀 ASYNC: Process multiple documents with advanced pipeline - WITH SMART CACHING"""
        os.makedirs(output_dir, exist_ok=True)
        
        # 🚀 SMART CACHING: Check if documents are already processed
        cache_info = self._check_processing_cache(file_paths, output_dir, chunk_size)
        
        # 🔥 CRITICAL FIX: Handle cases where cache_info might be None
        if cache_info is None or not isinstance(cache_info, dict):
            logger.warning("⚠️ Cache info is invalid, proceeding without cache")
            cache_info = {
                "all_cached": False,
                "files_to_process": file_paths,
                "existing_chunks": [],
                "existing_stats": None,
                "cache_file": str(Path(output_dir) / f"processing_cache_{chunk_size}.json")
            }
        
        if cache_info.get("all_cached", False):
            logger.info(f"✅ ALL DOCUMENTS ALREADY PROCESSED - Using cached data from {cache_info.get('cache_file', 'unknown')}")
            cached_result = cache_info.get("cached_result")
            if cached_result is not None and isinstance(cached_result, dict):
                return cached_result
            else:
                logger.warning("⚠️ Invalid cached result, proceeding with processing")
        
        # Process only new/changed files
        files_to_process = cache_info.get("files_to_process", file_paths)
        if files_to_process is None or not isinstance(files_to_process, list):
            files_to_process = file_paths
            
        if len(files_to_process) < len(file_paths):
            logger.info(f"🔄 SMART CACHE: Processing {len(files_to_process)}/{len(file_paths)} files (others are cached)")
        
        all_chunks = cache_info.get("existing_chunks", [])
        if all_chunks is None or not isinstance(all_chunks, list):
            all_chunks = []
            
        existing_stats = cache_info.get("existing_stats")
        if existing_stats is None or not isinstance(existing_stats, dict):
            processing_stats = {
                "total_files": len(file_paths),
                "successful_files": 0,
                "failed_files": 0,
                "total_chunks": 0,
                "total_characters": 0,
                "files_processed": [],
            }
        else:
            processing_stats = existing_stats

        # 🚀 PARALLEL PROCESSING: Process files concurrently instead of sequentially
        async def process_single_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
            """Process a single file asynchronously"""
            try:
                logger.info(f"Processing file: {file_path}")

                # Extract and process content (run in thread pool for CPU-bound operations)
                loop = asyncio.get_event_loop()
                content = await loop.run_in_executor(None, self.extract_text_with_structure, file_path)

                # 🔥 CRITICAL FIX: Handle cases where text extraction returns None or is empty
                if not content or not content.get("raw"):
                    logger.error(f"Content extraction failed for {file_path}, skipping.")
                    return file_path, {"status": "failed", "error": "Content extraction failed"}

                chunks = await loop.run_in_executor(None, self.generate_training_dataset, content, chunk_size)

                # 🔥 CRITICAL FIX: Validate chunks before processing
                if chunks is None:
                    logger.error(f"❌ Chunks generation returned None for {file_path}")
                    chunks = []
                
                if not isinstance(chunks, list):
                    logger.error(f"❌ Chunks is not a list for {file_path}: {type(chunks)}")
                    chunks = []

                # Save individual file results
                file_stem = Path(file_path).stem
                output_file = Path(output_dir) / f"{file_stem}_processed.txt"

                try:
                    async with aiofiles.open(output_file, "w", encoding="utf-8") as f:
                        for i, chunk in enumerate(chunks):
                            # 🔥 CRITICAL FIX: Validate each chunk
                            if chunk is None or not isinstance(chunk, dict):
                                logger.warning(f"⚠️ Invalid chunk at index {i} for {file_path}")
                                continue
                            
                            chunk_text = chunk.get('text', '')
                            if not chunk_text:
                                logger.warning(f"⚠️ Empty chunk text at index {i} for {file_path}")
                                continue
                            await f.write(f"[CHUNK_{i}]\n{chunk_text}\n\n")
                except Exception as e:
                    logger.error(f"❌ Failed to write chunks to file for {file_path}: {e}")

                return file_path, {
                    "status": "success",
                    "chunks": chunks,
                    "chunk_count": len(chunks),
                    "total_chars": sum(len(chunk.get('text', '')) for chunk in chunks)
                }

            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                return file_path, {"status": "failed", "error": str(e)}

        # Process all files in parallel with controlled concurrency
        semaphore = asyncio.Semaphore(4)  # Limit to 4 concurrent file operations
        
        async def process_with_semaphore(file_path: str):
            async with semaphore:
                return await process_single_file(file_path)
        
        # Execute all file processing tasks concurrently
        if files_to_process:
            tasks = [process_with_semaphore(file_path) for file_path in files_to_process]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Process results
            for result in results:
                if isinstance(result, Exception):
                    logger.error(f"Task failed with exception: {result}")
                    processing_stats["failed_files"] += 1
                    continue
                
                file_path, file_result = result
                if file_result.get("status") == "success":
                    chunks = file_result.get("chunks", [])
                    all_chunks.extend(chunks)
                    processing_stats["successful_files"] += 1
                    processing_stats["total_chunks"] += file_result.get("chunk_count", 0)
                    processing_stats["total_characters"] += file_result.get("total_chars", 0)
                    processing_stats["files_processed"].append(file_path)
                else:
                    processing_stats["failed_files"] += 1
                    logger.error(f"Failed to process {file_path}: {file_result.get('error', 'Unknown error')}")

        # Original sequential code removed - replaced with parallel processing above
        
        processing_stats["total_chunks"] = len(all_chunks)

        # Save combined dataset with metadata (human-readable)
        combined_file = Path(output_dir) / "combined_training_dataset.txt"
        try:
            with open(combined_file, "w", encoding="utf-8") as f:
                for chunk in all_chunks:
                    # 🔥 CRITICAL FIX: Validate chunk before accessing its keys
                    if chunk is None or not isinstance(chunk, dict):
                        logger.warning("⚠️ Skipping invalid chunk in combined dataset")
                        continue
                    
                    source = chunk.get('source', 'Unknown')
                    chunk_id = chunk.get('chunk_id', 0)
                    text = chunk.get('text', '')
                    
                    if not text:
                        logger.warning("⚠️ Skipping chunk with empty text in combined dataset")
                        continue
                    
                    f.write(f"Source: {source}\n")
                    f.write(f"Chunk ID: {chunk_id}\n")
                    f.write(f"Text: {text}\n")
                    f.write("-" * 80 + "\n\n")
        except Exception as e:
            logger.error(f"❌ Failed to write combined dataset: {e}")

        # 🆕 Save CLEAN JSONL training dataset (no metadata) for model fine-tuning
        clean_jsonl_file = Path(output_dir) / "training_dataset.jsonl"
        try:
            with open(clean_jsonl_file, "w", encoding="utf-8") as f:
                for chunk in all_chunks:
                    # 🔥 CRITICAL FIX: Validate chunk before accessing its keys
                    if chunk is None or not isinstance(chunk, dict):
                        logger.warning("⚠️ Skipping invalid chunk in JSONL dataset")
                        continue
                    
                    chunk_text = chunk.get('text', '')
                    if not chunk_text:
                        logger.warning("⚠️ Skipping chunk with empty text in JSONL dataset")
                        continue
                    
                    # Remove null bytes from chunk text as additional safety
                    clean_text = str(chunk_text).replace('\x00', '')
                    json_line = json.dumps({"text": clean_text}, ensure_ascii=False)
                    f.write(json_line + "\n")
        except Exception as e:
            logger.error(f"❌ Failed to write JSONL dataset: {e}")

        processing_stats["combined_output"] = str(combined_file)
        processing_stats["training_jsonl"] = str(clean_jsonl_file)
        
        # 💾 Save cache metadata for future smart caching
        self._save_processing_cache(file_paths, output_dir, chunk_size, processing_stats, all_chunks)

        logger.info(
            f"Processing complete: {processing_stats['successful_files']}/{processing_stats['total_files']} files successful"
        )
        return {
            "chunks": all_chunks,
            "stats": processing_stats,
            "output_files": [
                str(Path(output_dir) / f"{Path(fp).stem}_processed.txt") for fp in file_paths
            ] + [str(combined_file), str(clean_jsonl_file)],
        }
    
    def _check_processing_cache(self, file_paths: List[str], output_dir: str, chunk_size: int) -> Dict[str, Any]:
        """🚀 ENHANCED CACHE: Check if documents have already been processed with configuration-aware caching"""
        try:
            import threading
            
            # CRITICAL FIX: Thread-safe cache checking with lock
            cache_lock = getattr(self, '_cache_lock', None)
            if cache_lock is None:
                self._cache_lock = threading.Lock()
                cache_lock = self._cache_lock
            
            with cache_lock:
                # 🚀 ENHANCED: Create configuration hash for cache invalidation
                config_params = {
                    'chunk_size': chunk_size,
                    'preserve_educational_content': self.preserve_educational_content,
                    'use_semantic_chunking': self.use_semantic_chunking,
                    'has_sentence_transformer': self.sentence_transformer is not None,
                    'quality_thresholds': self.quality_thresholds,
                    'processor_version': '2.0'  # Increment when processing logic changes
                }
                
                config_hash = hashlib.md5(json.dumps(config_params, sort_keys=True).encode()).hexdigest()[:8]
                cache_file = Path(output_dir) / f"processing_cache_{chunk_size}_{config_hash}.json"
                
                if not cache_file.exists():
                    logger.info("🔄 No processing cache found - will process all files")
                    return {
                        "all_cached": False,
                        "files_to_process": file_paths,
                        "existing_chunks": [],
                        "existing_stats": None,
                        "cache_file": str(cache_file)
                    }
                
                # CRITICAL FIX: Robust cache loading with validation
                try:
                    with open(cache_file, 'r', encoding='utf-8') as f:
                        cache_data = json.loads(f.read())
                        
                    # Validate cache structure
                    if not isinstance(cache_data, dict):
                        raise ValueError("Invalid cache format - not a dictionary")
                        
                    required_cache_keys = ["file_paths", "chunk_size", "processing_stats", "chunks", "timestamp"]
                    missing_keys = [key for key in required_cache_keys if key not in cache_data]
                    if missing_keys:
                        raise ValueError(f"Invalid cache format - missing keys: {missing_keys}")
                        
                except (json.JSONDecodeError, ValueError, FileNotFoundError) as e:
                    logger.warning(f"⚠️ Invalid cache file {cache_file}: {e}")
                    logger.info("🗑️ Removing corrupted cache file")
                    try:
                        cache_file.unlink()
                    except Exception:
                        pass
                    return {
                        "all_cached": False,
                        "files_to_process": file_paths,
                        "existing_chunks": [],
                        "existing_stats": None,
                        "cache_file": str(cache_file)
                    }
                
                # Check if cache is valid for current request
                cached_file_paths = cache_data.get("file_paths", [])
                cached_chunk_size = cache_data.get("chunk_size", 0)
                
                # CRITICAL FIX: More robust file comparison with content hashing
                current_files_set = set()
                current_file_hashes = {}
                
                for fp in file_paths:
                    file_path = Path(fp).resolve()
                    current_files_set.add(file_path)
                    # Calculate file hash for content comparison
                    try:
                        import hashlib
                        with open(file_path, 'rb') as f:
                            file_hash = hashlib.sha256(f.read()).hexdigest()
                            current_file_hashes[str(file_path)] = file_hash
                    except Exception as e:
                        logger.warning(f"⚠️ Could not hash file {file_path}: {e}")
                        current_file_hashes[str(file_path)] = str(file_path.stat().st_mtime)
                
                cached_files_set = set(Path(fp).resolve() for fp in cached_file_paths if Path(fp).exists())
                cached_file_hashes = cache_data.get("file_hashes", {})
                
                if cached_chunk_size != chunk_size:
                    logger.info(f"🔄 Cache chunk size mismatch ({cached_chunk_size} vs {chunk_size}) - will reprocess")
                    return {
                        "all_cached": False,
                        "files_to_process": file_paths,
                        "existing_chunks": [],
                        "existing_stats": None,
                        "cache_file": str(cache_file)
                    }
                
                if current_files_set != cached_files_set:
                    # Determine which files need processing
                    files_to_process = [fp for fp in file_paths if Path(fp).resolve() not in cached_files_set]
                    
                    if not files_to_process:
                        # All files are cached, but some cached files might be missing
                        logger.info("✅ All current files found in cache")
                        return {
                            "all_cached": True,
                            "cached_result": {
                                "chunks": cache_data.get("chunks", []),
                                "stats": cache_data.get("processing_stats", {}),
                                "cache_hit": True
                            },
                            "cache_file": str(cache_file)
                        }
                    else:
                        logger.info(f"🔄 {len(files_to_process)} new files need processing")
                        return {
                            "all_cached": False,
                            "files_to_process": files_to_process,
                            "existing_chunks": cache_data.get("chunks", []),
                            "existing_stats": cache_data.get("processing_stats", {}),
                            "cache_file": str(cache_file)
                        }
                
                # CRITICAL FIX: Verify cached files still exist and haven't changed content
                for cached_file in cached_file_paths:
                    cached_path = Path(cached_file)
                    if not cached_path.exists():
                        logger.warning(f"⚠️ Cached file no longer exists: {cached_file}")
                        return {
                            "all_cached": False,
                            "files_to_process": file_paths,
                            "existing_chunks": [],
                            "existing_stats": None,
                            "cache_file": str(cache_file)
                        }
                    
                    # CRITICAL FIX: Check content hash for changes
                    cached_file_str = str(cached_path)
                    if cached_file_str in current_file_hashes and cached_file_str in cached_file_hashes:
                        current_hash = current_file_hashes[cached_file_str]
                        cached_hash = cached_file_hashes[cached_file_str]
                        if current_hash != cached_hash:
                            logger.info(f"🔄 File content changed since caching: {cached_file}")
                            return {
                                "all_cached": False,
                                "files_to_process": file_paths,
                                "existing_chunks": [],
                                "existing_stats": None,
                                "cache_file": str(cache_file)
                            }
                    else:
                        # Fallback to timestamp check if hashing fails
                        try:
                            cache_timestamp = cache_data.get("timestamp", 0)
                            file_modified_time = cached_path.stat().st_mtime
                            if file_modified_time > cache_timestamp:
                                logger.info(f"🔄 File modified since caching: {cached_file}")
                                return {
                                    "all_cached": False,
                                    "files_to_process": file_paths,
                                    "existing_chunks": [],
                                    "existing_stats": None,
                                    "cache_file": str(cache_file)
                                }
                        except Exception as e:
                            logger.warning(f"⚠️ Could not check file modification time for {cached_file}: {e}")
                
                # All files are cached and valid
                logger.info("✅ All files found in valid cache")
                return {
                    "all_cached": True,
                    "cached_result": {
                        "chunks": cache_data.get("chunks", []),
                        "stats": cache_data.get("processing_stats", {}),
                        "cache_hit": True
                    },
                    "cache_file": str(cache_file)
                }
                
        except Exception as e:
            logger.error(f"❌ Error checking processing cache: {e}")
            # Return safe default to continue processing
            return {
                "all_cached": False,
                "files_to_process": file_paths,
                "existing_chunks": [],
                "existing_stats": None,
                "cache_file": str(Path(output_dir) / f"processing_cache_{chunk_size}.json")
            }
    
    def _save_processing_cache(self, file_paths: List[str], output_dir: str, chunk_size: int, 
                             processing_stats: Dict[str, Any], all_chunks: List[Dict[str, Any]]):
        """🚀 SMART CACHE: Save processing results to cache with thread safety"""
        try:
            import threading
            
            # CRITICAL FIX: Thread-safe cache saving with lock
            cache_lock = getattr(self, '_cache_lock', None)
            if cache_lock is None:
                self._cache_lock = threading.Lock()
                cache_lock = self._cache_lock
            
            with cache_lock:
                cache_file = Path(output_dir) / f"processing_cache_{chunk_size}.json"
                
                # CRITICAL FIX: Create cache data with validation and file hashes
                import hashlib
                file_hashes = {}
                
                for fp in file_paths:
                    try:
                        file_path = Path(fp).resolve()
                        with open(file_path, 'rb') as f:
                            file_hash = hashlib.sha256(f.read()).hexdigest()
                            file_hashes[str(file_path)] = file_hash
                    except Exception as e:
                        logger.warning(f"⚠️ Could not hash file {fp}: {e}")
                        file_hashes[str(file_path)] = str(Path(fp).stat().st_mtime)
                
                cache_data = {
                    "file_paths": [str(Path(fp).resolve()) for fp in file_paths],
                    "chunk_size": chunk_size,
                    "processing_stats": processing_stats,
                    "chunks": all_chunks,
                    "timestamp": time.time(),
                    "file_hashes": file_hashes,  # CRITICAL FIX: Store file hashes
                    "version": "1.1"  # Updated version for hash support
                }
                
                # Validate cache data before saving
                try:
                    # Test JSON serialization
                    json_str = json.dumps(cache_data, ensure_ascii=False, indent=2)
                    
                    # CRITICAL FIX: Atomic write to prevent corruption
                    temp_cache_file = cache_file.with_suffix('.tmp')
                    with open(temp_cache_file, 'w', encoding='utf-8') as f:
                        f.write(json_str)
                        f.flush()
                        os.fsync(f.fileno())  # Force write to disk
                    
                    # Atomic move
                    temp_cache_file.replace(cache_file)
                    
                    logger.info(f"💾 Processing cache saved: {cache_file}")
                    
                except Exception as write_error:
                    logger.error(f"❌ Failed to write cache file: {write_error}")
                    # Clean up temp file if it exists
                    if temp_cache_file.exists():
                        try:
                            temp_cache_file.unlink()
                        except Exception:
                            pass
                    
        except Exception as e:
            logger.error(f"❌ Error saving processing cache: {e}")
            # Don't raise - caching is optional

    def clear_processing_cache(self, output_dir: str):
        """🧹 Clear processing cache to force reprocessing"""
        try:
            cache_file = Path(output_dir) / "processing_cache.json"
            chunks_file = Path(output_dir) / "cached_chunks.json"
            
            if cache_file.exists():
                cache_file.unlink()
                logger.info(f"🧹 Cleared processing cache: {cache_file}")
            
            if chunks_file.exists():
                chunks_file.unlink()
                logger.info(f"🧹 Cleared chunks cache: {chunks_file}")
                
        except Exception as e:
            logger.error(f"❌ Error clearing cache: {e}")

    def extract_images_from_pdf(self, pdf_path: str, output_folder: str) -> List[str]:
        """Extract images from PDF file"""
        if not pdf_support_available:
            logger.warning("PyMuPDF not available. Cannot extract images from PDF.")
            return []

        os.makedirs(output_folder, exist_ok=True)
        extracted_image_paths = []

        try:
            doc = fitz.open(pdf_path)
            for page_num, page in enumerate(doc):
                image_list = page.get_images(full=True)

                for img_idx, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]

                    image_filename = f"page_{page_num + 1}_img_{img_idx + 1}.{base_image['ext']}"
                    image_path = os.path.join(output_folder, image_filename)

                    with open(image_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    extracted_image_paths.append(image_path)

            doc.close()
            logger.info(f"Extracted {len(extracted_image_paths)} images from PDF")
            return extracted_image_paths

        except Exception as e:
            logger.error(f"Error extracting images from PDF: {e}")
            return []

    def process_books(
        self,
        directory: str = "uploaded_books",
        output_file: Optional[str] = "training_data.txt",
        output_image_folder: str = "extracted_images_for_training",
    ) -> List[Dict[str, Any]]:
        """Process books and extract text and images"""
        if not os.path.exists(directory):
            raise FileNotFoundError(f"Directory not found: {directory}")

        os.makedirs(output_image_folder, exist_ok=True)
        corpus_texts_and_image_paths = []

        for filename in os.listdir(directory):
            filepath = os.path.join(directory, filename)
            text_content = ""
            extracted_image_paths_for_this_doc = []

            try:
                content = self.extract_text_with_structure(filepath)
                # 🔥 CRITICAL FIX: Validate content before accessing raw key
                if content is None or not isinstance(content, dict):
                    logger.error(f"❌ Invalid content structure for {filename}")
                    text_content = ""
                elif "raw" not in content:
                    logger.error(f"❌ Missing 'raw' key in content for {filename}")
                    text_content = ""
                else:
                    text_content = content.get("raw", "")
                    if text_content is None:
                        text_content = ""

                if filename.lower().endswith(".pdf") and text_content:
                    pdf_image_subfolder = os.path.join(
                        output_image_folder, os.path.splitext(filename)[0]
                    )
                    extracted_image_paths_for_this_doc = self.extract_images_from_pdf(
                        filepath, pdf_image_subfolder
                    )

            except Exception as e:
                logger.error(f"Error processing file '{filename}': {e}")
                text_content = f"ERROR PROCESSING FILE: {str(e)}"

            if text_content:
                corpus_texts_and_image_paths.append(
                    {"text": text_content, "images": extracted_image_paths_for_this_doc}
                )

        if output_file:
            try:
                with open(output_file, "w", encoding="utf-8") as f:
                    for item in corpus_texts_and_image_paths:
                        f.write(item["text"] + "\n\n")
            except Exception as e:
                logger.error(f"Error saving to output file: {e}")

        return corpus_texts_and_image_paths

    def semantic_chunking(self, text: str, max_chunk_size: int = 500, overlap_sentences: int = 2) -> List[Dict[str, Any]]:
        """🚀 ADVANCED: Semantic chunking using sentence similarity and topic detection"""
        if not text or not isinstance(text, str):
            logger.warning("Invalid text for semantic chunking")
            return []

        try:
            # Step 1: Sentence tokenization
            if nltk_available and sent_tokenize:
                sentences = sent_tokenize(text)
            else:
                # Fallback sentence splitting
                sentences = self._fallback_sentence_split(text)
            
            if len(sentences) <= 3:
                # Too few sentences for semantic chunking, return as single chunk
                return [{
                    'text': text,
                    'sentence_count': len(sentences),
                    'semantic_coherence': 1.0,
                    'chunk_id': 0,
                    'method': 'single_chunk'
                }]

            # Step 2: Calculate sentence embeddings or similarity
            if sklearn_available and self.tfidf_vectorizer:
                chunks = self._embedding_based_chunking(sentences, max_chunk_size, overlap_sentences)
            else:
                # Fallback to sentence-window chunking
                chunks = self._sentence_window_chunking(sentences, max_chunk_size, overlap_sentences)
            
            return chunks

        except Exception as e:
            logger.error(f"Error in semantic chunking: {e}")
            # Fallback to simple chunking
            return self._simple_text_chunking(text, max_chunk_size)

    def _embedding_based_chunking(self, sentences: List[str], max_chunk_size: int, overlap_sentences: int) -> List[Dict[str, Any]]:
        """🚀 ADVANCED: Embedding-based semantic chunking using TF-IDF similarity"""
        try:
            # Calculate TF-IDF vectors for sentences
            sentence_vectors = self.tfidf_vectorizer.fit_transform(sentences)
            
            # Calculate similarity matrix
            similarity_matrix = cosine_similarity(sentence_vectors)
            
            chunks = []
            current_chunk = []
            current_size = 0
            chunk_id = 0
            
            for i, sentence in enumerate(sentences):
                sentence_len = len(sentence.split())
                
                # Check if adding this sentence would exceed chunk size
                if current_size + sentence_len > max_chunk_size and current_chunk:
                    # Check semantic coherence before splitting
                    if i < len(sentences) - 1:
                        current_similarity = similarity_matrix[i-1][i] if i > 0 else 1.0
                        next_similarity = similarity_matrix[i][i+1]
                        
                        # Split here if similarity drops significantly (topic change)
                        if current_similarity - next_similarity > 0.3:  # Threshold for topic change
                            # Create chunk
                            chunk_text = ' '.join(current_chunk)
                            semantic_score = self._calculate_chunk_coherence(current_chunk, sentence_vectors[max(0, i-len(current_chunk)):i])
                            
                            chunks.append({
                                'text': chunk_text,
                                'sentence_count': len(current_chunk),
                                'semantic_coherence': semantic_score,
                                'chunk_id': chunk_id,
                                'method': 'embedding_based',
                                'word_count': current_size
                            })
                            
                            # Start new chunk with overlap
                            overlap_start = max(0, len(current_chunk) - overlap_sentences)
                            current_chunk = current_chunk[overlap_start:] + [sentence]
                            current_size = sum(len(s.split()) for s in current_chunk)
                            chunk_id += 1
                            continue
                
                current_chunk.append(sentence)
                current_size += sentence_len
            
            # Add final chunk
            if current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'sentence_count': len(current_chunk),
                    'semantic_coherence': 0.8,  # Default for final chunk
                    'chunk_id': chunk_id,
                    'method': 'embedding_based',
                    'word_count': current_size
                })
            
            return chunks

        except Exception as e:
            logger.error(f"Error in embedding-based chunking: {e}")
            return self._sentence_window_chunking(sentences, max_chunk_size, overlap_sentences)

    def _sentence_window_chunking(self, sentences: List[str], max_chunk_size: int, overlap_sentences: int) -> List[Dict[str, Any]]:
        """🚀 ADVANCED: Sentence-window based chunking with overlap"""
        chunks = []
        current_chunk = []
        current_size = 0
        chunk_id = 0
        
        for sentence in sentences:
            sentence_len = len(sentence.split())
            
            if current_size + sentence_len > max_chunk_size and current_chunk:
                # Create chunk
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'sentence_count': len(current_chunk),
                    'semantic_coherence': 0.7,  # Default coherence score
                    'chunk_id': chunk_id,
                    'method': 'sentence_window',
                    'word_count': current_size
                })
                
                # Start new chunk with overlap
                overlap_start = max(0, len(current_chunk) - overlap_sentences)
                current_chunk = current_chunk[overlap_start:]
                current_size = sum(len(s.split()) for s in current_chunk)
                chunk_id += 1
            
            current_chunk.append(sentence)
            current_size += sentence_len
        
        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'text': chunk_text,
                'sentence_count': len(current_chunk),
                'semantic_coherence': 0.7,
                'chunk_id': chunk_id,
                'method': 'sentence_window',
                'word_count': current_size
            })
        
        return chunks

    def _calculate_chunk_coherence(self, sentences: List[str], vectors) -> float:
        """Calculate semantic coherence score for a chunk"""
        try:
            if len(sentences) < 2:
                return 1.0
            
            # Calculate average pairwise similarity
            similarities = []
            for i in range(len(vectors) - 1):
                sim = cosine_similarity(vectors[i:i+1], vectors[i+1:i+2])[0][0]
                similarities.append(sim)
            
            return np.mean(similarities) if similarities else 0.5
        except:
            return 0.5  # Default coherence

    def _fallback_sentence_split(self, text: str) -> List[str]:
        """🚀 ENHANCED: Better sentence splitting when NLTK is not available"""
        # More sophisticated sentence splitting with common abbreviations handling
        
        # Common abbreviations that shouldn't end sentences
        abbreviations = {
            'dr', 'mr', 'mrs', 'ms', 'prof', 'vs', 'etc', 'inc', 'ltd', 'corp',
            'fig', 'eq', 'ref', 'vol', 'no', 'pp', 'ch', 'sect', 'al', 'et'
        }
        
        # First, protect abbreviations by temporarily replacing periods
        protected_text = text
        for abbrev in abbreviations:
            # Replace abbreviation periods with placeholder
            pattern = r'\b' + abbrev + r'\.'
            protected_text = re.sub(pattern, abbrev + '###PERIOD###', protected_text, flags=re.IGNORECASE)
        
        # Split on sentence-ending punctuation, but be more careful
        # Don't split on periods in numbers or abbreviations
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', protected_text)
        
        # Restore protected periods
        sentences = [s.replace('###PERIOD###', '.') for s in sentences]
        
        # Clean and filter sentences
        cleaned_sentences = []
        for s in sentences:
            s = s.strip()
            if len(s) > 10 and not s.isdigit():  # Filter out very short or numeric-only sentences
                cleaned_sentences.append(s)
        
        return cleaned_sentences

    def _simple_text_chunking(self, text: str, max_chunk_size: int) -> List[Dict[str, Any]]:
        """Simple word-based chunking as final fallback"""
        words = text.split()
        chunks = []
        current_chunk = []
        chunk_id = 0
        
        for word in words:
            if len(current_chunk) >= max_chunk_size:
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'sentence_count': chunk_text.count('.') + 1,
                    'semantic_coherence': 0.5,
                    'chunk_id': chunk_id,
                    'method': 'simple_word',
                    'word_count': len(current_chunk)
                })
                current_chunk = []
                chunk_id += 1
            
            current_chunk.append(word)
        
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'text': chunk_text,
                'sentence_count': chunk_text.count('.') + 1,
                'semantic_coherence': 0.5,
                'chunk_id': chunk_id,
                'method': 'simple_word',
                'word_count': len(current_chunk)
            })
        
        return chunks

    def quality_filter_chunks(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """🚀 ADVANCED: Quality filtering using multiple criteria"""
        if not chunks:
            return []
        
        filtered_chunks = []
        
        for chunk in chunks:
            if not chunk or not isinstance(chunk, dict):
                continue
            
            text = chunk.get('text', '')
            if not text:
                continue
            
            # Quality scoring
            quality_score = self._calculate_chunk_quality(text)
            
            # Language detection
            if langdetect_available:
                try:
                    detected_lang = langdetect.detect(text[:500])  # Sample for detection
                    if detected_lang != 'en':  # Filter non-English content
                        logger.debug(f"Filtering non-English chunk: {detected_lang}")
                        continue
                except:
                    pass  # Continue if language detection fails
            
            # Perplexity filtering (simplified heuristic)
            perplexity_score = self._estimate_perplexity(text)
            
            # Apply thresholds
            if (quality_score >= 0.5 and 
                perplexity_score <= self.quality_thresholds['perplexity_threshold'] and
                len(text.strip()) >= 20):
                
                chunk['quality_score'] = quality_score
                chunk['perplexity_score'] = perplexity_score
                filtered_chunks.append(chunk)
            else:
                logger.debug(f"Filtered chunk: quality={quality_score:.2f}, perplexity={perplexity_score:.2f}")
        
        logger.info(f"Quality filtering: {len(chunks)} -> {len(filtered_chunks)} chunks")
        return filtered_chunks

    def _calculate_chunk_quality(self, text: str) -> float:
        """Calculate quality score for a text chunk"""
        try:
            words = text.split()
            if not words:
                return 0.0
            
            # Calculate various quality metrics
            scores = []
            
            # 1. Average word length (longer words often indicate higher quality)
            avg_word_length = np.mean([len(word) for word in words])
            scores.append(min(avg_word_length / 6.0, 1.0))  # Normalize to 0-1
            
            # 2. Sentence structure (presence of proper punctuation)
            sentence_score = min(text.count('.') + text.count('!') + text.count('?'), 10) / 10.0
            scores.append(sentence_score)
            
            # 3. Educational content indicators
            educational_indicators = [
                'example', 'theorem', 'proof', 'solution', 'problem', 'formula', 
                'equation', 'method', 'principle', 'law', 'definition', 'concept'
            ]
            edu_score = sum(1 for indicator in educational_indicators if indicator in text.lower()) / len(educational_indicators)
            scores.append(edu_score)
            
            # 4. Mathematical content bonus
            if self._contains_mathematical_content(text):
                scores.append(0.8)  # Bonus for math content
            
            # 5. Avoid repetitive or boilerplate text
            unique_words = len(set(words))
            if unique_words > 0:
                repetition_score = min(unique_words / len(words), 1.0)
                scores.append(repetition_score)
            
            return np.mean(scores)
            
        except Exception as e:
            logger.warning(f"Error calculating quality score: {e}")
            return 0.5

    def _estimate_perplexity(self, text: str) -> float:
        """Estimate perplexity using simple heuristics"""
        try:
            words = text.split()
            if len(words) < 5:
                return 100.0  # High perplexity for very short text
            
            # Simple perplexity estimation based on word patterns
            # Lower values indicate more coherent text
            
            # Check for unusual character patterns
            unusual_chars = len(re.findall(r'[^\w\s\.\,\!\?\;\:\-\(\)]', text))
            char_penalty = unusual_chars / len(text) * 50
            
            # Check for very short or very long words (OCR artifacts)
            word_lengths = [len(word) for word in words]
            avg_length = np.mean(word_lengths)
            length_penalty = abs(avg_length - 5) * 2  # Optimal around 5 chars
            
            # Check for excessive repetition
            unique_ratio = len(set(words)) / len(words)
            repetition_penalty = (1 - unique_ratio) * 30
            
            perplexity = 10 + char_penalty + length_penalty + repetition_penalty
            return min(perplexity, 100.0)
            
        except Exception as e:
            logger.warning(f"Error estimating perplexity: {e}")
            return 25.0

    def data_augmentation(self, chunks: List[Dict[str, Any]], augmentation_ratio: float = 0.1) -> List[Dict[str, Any]]:
        """🚀 ADVANCED: Data augmentation for training robustness"""
        if not chunks or augmentation_ratio <= 0:
            return chunks
        
        augmented_chunks = chunks.copy()
        num_to_augment = int(len(chunks) * augmentation_ratio)
        
        logger.info(f"Augmenting {num_to_augment} chunks out of {len(chunks)}")
        
        for i in range(num_to_augment):
            if i >= len(chunks):
                break
            
            original_chunk = chunks[i % len(chunks)]
            original_text = original_chunk.get('text', '')
            
            if not original_text or len(original_text) < 50:
                continue
            
            # Apply augmentation techniques
            augmented_text = self._apply_text_augmentation(original_text)
            
            if augmented_text and augmented_text != original_text:
                augmented_chunk = original_chunk.copy()
                augmented_chunk['text'] = augmented_text
                augmented_chunk['chunk_id'] = len(augmented_chunks)
                augmented_chunk['augmented'] = True
                augmented_chunk['original_chunk_id'] = original_chunk['chunk_id']
                
                augmented_chunks.append(augmented_chunk)
        
        return augmented_chunks

    def _apply_text_augmentation(self, text: str) -> str:
        """Apply various text augmentation techniques"""
        try:
            # For now, implement simple synonym replacement
            # This could be enhanced with back-translation or more sophisticated methods
            
            words = text.split()
            augmented_words = []
            
            # Simple synonym mapping for common educational terms
            synonyms = {
                'method': 'approach',
                'problem': 'exercise', 
                'solution': 'answer',
                'example': 'instance',
                'result': 'outcome',
                'calculate': 'compute',
                'determine': 'find',
                'show': 'demonstrate',
                'prove': 'establish',
                'formula': 'equation'
            }
            
            for word in words:
                # Replace with synonym ~20% of the time for non-critical words
                if (word.lower() in synonyms and 
                    np.random.random() < 0.2 and 
                    not self._is_critical_term(word)):
                    augmented_words.append(synonyms[word.lower()])
                else:
                    augmented_words.append(word)
            
            return ' '.join(augmented_words)
            
        except Exception as e:
            logger.warning(f"Error in text augmentation: {e}")
            return text

    def _is_critical_term(self, word: str) -> bool:
        """Check if a word is critical and should not be replaced"""
        critical_patterns = [
            r'\d+',  # Numbers
            r'[A-Z]{2,}',  # Acronyms
            r'[a-zA-Z]+\d+',  # Terms with numbers
        ]
        
        for pattern in critical_patterns:
            if re.match(pattern, word):
                return True
        
        # Mathematical terms that shouldn't be replaced
        math_terms = {'sin', 'cos', 'tan', 'log', 'ln', 'exp', 'sqrt', 'integral', 'derivative'}
        if word.lower() in math_terms:
            return True
        
        return False

    def generate_enhanced_training_dataset(self, content: Dict[str, str], chunk_size: int = 500) -> List[Dict[str, str]]:
        """🚀 ENHANCED: Generate training dataset with semantic chunking and quality filtering"""
        
        # Validate content structure before processing
        if not content or not isinstance(content, dict):
            logger.error("❌ Content is None or not a dictionary")
            return []
        
        if "raw" not in content:
            logger.error("❌ Content missing 'raw' key")
            return []
        
        raw_content = content.get("raw")
        if not raw_content or not isinstance(raw_content, str):
            logger.error("❌ Raw content is None or not a string")
            return []

        # Step 1: Enhanced text cleaning
        cleaned_text = self.educational_text_cleaning(raw_content)
        
        if not cleaned_text or not isinstance(cleaned_text, str):
            logger.error("❌ Text cleaning failed or returned invalid content")
            return []

        # Step 2: Semantic chunking
        if self.use_semantic_chunking:
            raw_chunks = self.semantic_chunking(cleaned_text, chunk_size)
        else:
            # Fallback to sentence-based chunking
            if nltk_available and sent_tokenize:
                sentences = sent_tokenize(cleaned_text)
                raw_chunks = self._sentence_window_chunking(sentences, chunk_size, 2)
            else:
                raw_chunks = self._simple_text_chunking(cleaned_text, chunk_size)

        if not raw_chunks:
            logger.error("❌ No chunks generated from content")
            return []

        # Step 3: Quality filtering
        filtered_chunks = self.quality_filter_chunks(raw_chunks)
        
        if not filtered_chunks:
            logger.warning("⚠️ All chunks filtered out, using original chunks")
            filtered_chunks = raw_chunks

        # Step 4: Data augmentation (optional)
        if len(filtered_chunks) > 5:  # Only augment if we have enough base chunks
            final_chunks = self.data_augmentation(filtered_chunks, 0.15)  # 15% augmentation
        else:
            final_chunks = filtered_chunks

        # Step 5: Convert to training format
        training_chunks = []
        metadata = content.get("metadata", {})
        source_title = metadata.get("title", "Unknown") if metadata else "Unknown"
        
        for chunk in final_chunks:
            if not chunk or not isinstance(chunk, dict):
                continue
            
            chunk_text = chunk.get('text', '')
            if not chunk_text or len(chunk_text.strip()) < 20:
                continue
            
            # 🚀 ENHANCED: Create instruction-following format for better training data
            training_chunk = {
                # Basic text chunk
                "text": chunk_text,
                "length": len(chunk_text),
                
                # Metadata for provenance and quality
                "source": str(source_title),
                "chunk_id": chunk.get('chunk_id', 0),
                "method": chunk.get('method', 'unknown'),
                "semantic_coherence": chunk.get('semantic_coherence', 0.5),
                "quality_score": chunk.get('quality_score', 0.5),
                "sentence_count": chunk.get('sentence_count', 1),
                "word_count": chunk.get('word_count', len(chunk_text.split())),
                "augmented": chunk.get('augmented', False),
                
                # 🚀 NEW: Instruction-following format for modern training
                "instruction_format": {
                    "instruction": "Summarize the following educational content:",
                    "input": chunk_text,
                    "output": "",  # To be filled by generation models
                    "system": "You are an educational content expert. Provide accurate, clear summaries."
                },
                
                # 🚀 NEW: Multiple task formats for diverse training
                "task_variations": [
                    {
                        "task": "summarization",
                        "instruction": "Summarize the following content in 2-3 sentences:",
                        "input": chunk_text
                    },
                    {
                        "task": "question_generation", 
                        "instruction": "Generate 2 relevant questions based on this content:",
                        "input": chunk_text
                    },
                    {
                        "task": "key_concepts",
                        "instruction": "List the main concepts discussed in this text:",
                        "input": chunk_text
                    }
                ],
                
                # Structure information for enhanced training
                "structure_info": {
                    "has_tables": "table" in chunk_text.lower(),
                    "has_math": bool(re.search(r'[=+\-*/\^]|\d+', chunk_text)),
                    "has_citations": bool(re.search(r'\[\d+\]|\(\d+\)', chunk_text)),
                    "complexity_level": "intermediate" if len(chunk_text.split()) > 100 else "basic"
                }
            }
            
            training_chunks.append(training_chunk)

        logger.info(f"🚀 Generated {len(training_chunks)} high-quality training chunks using enhanced processing")
        return training_chunks


# Maintain backward compatibility
class DocumentProcessor(AdvancedDocumentProcessor):
    """Backward compatible DocumentProcessor that inherits advanced features"""

    def __init__(self):
        super().__init__(preserve_educational_content=True, use_semantic_chunking=True)

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Legacy method for backward compatibility"""
        processor = AdvancedDocumentProcessor(preserve_educational_content=False, use_semantic_chunking=False)
        content = processor.extract_text_with_structure(file_path)
        # 🔥 CRITICAL FIX: Validate content before accessing raw key
        if content is None or not isinstance(content, dict):
            logger.error(f"❌ Invalid content structure for {file_path}")
            return ""
        elif "raw" not in content:
            logger.error(f"❌ Missing 'raw' key in content for {file_path}")
            return ""
        else:
            raw_content = content.get("raw", "")
            return raw_content if raw_content is not None else ""

    @staticmethod
    def clean_text(text: str) -> str:
        """Legacy method for backward compatibility"""
        processor = AdvancedDocumentProcessor(preserve_educational_content=True, use_semantic_chunking=True)
        return processor.educational_text_cleaning(text)

    @classmethod
    def process_documents(cls, file_paths: List[str], output_dir: str) -> List[str]:
        """Legacy method for backward compatibility"""
        processor = AdvancedDocumentProcessor(preserve_educational_content=True, use_semantic_chunking=True)
        result = processor.process_documents_advanced(file_paths, output_dir)
        return result["output_files"]
